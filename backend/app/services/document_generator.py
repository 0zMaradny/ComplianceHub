import os
import re
import logging

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.utils import sanitize_filename
from app.config import DEFAULT_LOGO_PATH
from app.services.client_config import get_client
from app.services.bilingual import t, TABLE_HEADERS, COVER_LABELS, METHODOLOGY, CONFIDENTIALITY

logger = logging.getLogger(__name__)

TUV_BLUE = RGBColor(0x00, 0x3D, 0x7A)
TUV_RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0xF2, 0xF2, 0xF2)

# ── RTL mode flag (set before calling any generator) ──────────────────────
_RTL_MODE = False


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number(doc):
    sect = doc.sections[0]
    footer = sect.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Page ')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    p.runs[-1]._r.addnext(fld_char1)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_char1.addnext(instr)
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    instr.addnext(fld_char2)


def set_col_widths(table, widths_pct, available_inches=6.5):
    available = Inches(available_inches)
    table.autofit = False
    for row in table.rows:
        try:
            cells = row.cells
        except Exception as e:
            logger.warning('set_col_widths: failed to access row cells: %s', e)
            continue
        for i, cell in enumerate(cells):
            if i < len(widths_pct):
                cell.width = int(available * widths_pct[i] / 100)


def set_landscape(doc):
    sect = doc.sections[0]
    sect.orientation = WD_ORIENT.LANDSCAPE
    sect.page_width, sect.page_height = sect.page_height, sect.page_width


def setup_document(doc, landscape=False, client_key=None):
    if landscape:
        set_landscape(doc)
    sect = doc.sections[0]
    sect.top_margin = Cm(2.0)
    sect.bottom_margin = Cm(2.0)
    sect.left_margin = Cm(2.5)
    sect.right_margin = Cm(2.5)

    header = sect.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ''
    run = hp.add_run('TÜV AUSTRIA')
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = TUV_BLUE
    run2 = hp.add_run('  |  Audit Documentation')
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_page_number(doc)


def add_header_row(table, headers, client_key=None):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, '003D7A')


def add_data_row(table, data, bold=False, color=None, client_key=None):
    row = table.add_row()
    for i, val in enumerate(data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        if i % 2 == 0:
            set_cell_shading(cell, 'F2F2F2')
    return row


def _parse_hex_color(val):
    if isinstance(val, RGBColor):
        return val
    if isinstance(val, str) and val.startswith('#'):
        return RGBColor(int(val[1:3], 16), int(val[3:5], 16), int(val[5:7], 16))
    return TUV_BLUE


def add_cover_page(doc, title, client_name, standard, date, lead_auditor='', client_key=None):
    # Resolve client for colors and language
    header_color = TUV_BLUE
    labels = COVER_LABELS.get("en", {})
    if client_key:
        c = get_client(client_key)
        if c:
            header_color = c.visual.primary_header
            lang = c.language
            labels = COVER_LABELS.get(lang, COVER_LABELS["en"])
    header_color = _parse_hex_color(header_color)

    if os.path.exists(DEFAULT_LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(DEFAULT_LOGO_PATH, width=Inches(3.0))

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = header_color

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TÜV AUSTRIA')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    info_items = [
        (labels.get("client", "Client"), client_name),
        (labels.get("standard", "Standard"), standard),
        (labels.get("date", "Date"), date),
    ]
    if lead_auditor:
        info_items.append((labels.get("lead_auditor", "Lead Auditor"), lead_auditor))

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}: ')
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = header_color
        run = p.add_run(value)
        run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('This document is confidential and intended solely for the use of the client.')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    doc.add_page_break()


def add_toc(doc, client_key=None):
    p = doc.add_paragraph()
    run = p.add_run('Table of Contents')
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = TUV_BLUE
    p.space_after = Pt(12)

    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r' TOC \o "1-3" \h \z \u '
    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    p2 = doc.add_paragraph()
    p2.space_after = Pt(4)
    for el in [fld_char_begin, instr, fld_char_separate, fld_char_end]:
        p2._p.append(el)

    doc.add_paragraph()


def add_section_heading(doc, text, client_key=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = TUV_BLUE
    p.space_after = Pt(6)
    p.space_before = Pt(12)
    return p


def add_sub_heading(doc, text, client_key=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = True
    p.space_after = Pt(4)
    return p


def add_body_text(doc, text, bold=False, client_key=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold
    p.space_after = Pt(4)
    return p


def add_bullet(doc, text, client_key=None):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_border_box(doc, client_key=None):
    sect = doc.sections[0]
    sect.left_margin = Cm(1.5)
    sect.right_margin = Cm(1.5)
    sect.top_margin = Cm(1.5)
    sect.bottom_margin = Cm(1.5)
    pg = sect._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for side, sz, color, space in [
        ('top', '24', '003D7A', '24'),
        ('left', '24', '003D7A', '24'),
        ('bottom', '24', '003D7A', '24'),
        ('right', '24', '003D7A', '24'),
    ]:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), sz)
        border.set(qn('w:color'), color)
        border.set(qn('w:space'), space)
        pgBorders.append(border)
    pg.append(pgBorders)


# ── RTL / Arabic text support ──────────────────────────────────────────────

def set_rtl_paragraph(paragraph):
    """Enable RTL for a paragraph (Arabic support)."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)


def add_body_text_rtl(doc, text, bold=False, client_key=None):
    """Add RTL body text (for Arabic)."""
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold
    p.space_after = Pt(4)
    return p


def _resolve_client(client_key: str = None):
    """Return (client, lang, rtl, header_color) tuple from client_key."""
    lang = "en"
    rtl = False
    header_color = TUV_BLUE
    client = None
    if client_key:
        client = get_client(client_key)
        if client:
            lang = client.language
            rtl = client.visual.rtl
            header_color = RGBColor(
                int(client.visual.primary_header[1:3], 16),
                int(client.visual.primary_header[3:5], 16),
                int(client.visual.primary_header[5:7], 16),
            )
    return client, lang, rtl, header_color


def add_cover_page_rtl(doc, title, client_name, standard, date, lead_auditor=''):
    """Add cover page with RTL support."""
    if os.path.exists(DEFAULT_LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(DEFAULT_LOGO_PATH, width=Inches(3.0))
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = TUV_BLUE
    doc.add_paragraph()
    info_items = [
        ('العميل / Client', client_name),
        ('المواصفة / Standard', standard),
        ('التاريخ / Date', date),
    ]
    if lead_auditor:
        info_items.append(('المراجع الرئيسي / Lead Auditor', lead_auditor))
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}: ')
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = TUV_BLUE
        run = p.add_run(value)
        run.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('هذه الوثيقة سرية ومخصصة لاستخدام العميل فقط')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    doc.add_page_break()


# ── Missing document generators ────────────────────────────────────────────

def generate_tnl(data, output_path, client_key: str = None):
    """Test/Nonconformity Log (TNL) document."""
    client, lang, rtl, header_color = _resolve_client(client_key)

    doc = Document()
    setup_document(doc, client_key=client_key)
    add_cover_page(doc, f'{TABLE_HEADERS[lang]["tnl"][0]} / TNL',
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'),
                   data.get('audit_date', 'N/A'),
                   client_key=client_key)

    add_section_heading(doc, t("nc_observation_log", lang), client_key=client_key)

    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    add_header_row(table, TABLE_HEADERS[lang]["tnl"], client_key=client_key)
    for entry in data.get('entries', []):
        etype = entry.get('type', '')
        row_color = None
        if etype == 'NC':
            row_color = TUV_RED
        elif etype == 'OFI':
            row_color = RGBColor(0xCC, 0x7A, 0x00)
        add_data_row(table, [
            entry.get('tnl_number', ''),
            entry.get('clause', ''),
            etype,
            entry.get('description', ''),
            entry.get('severity', ''),
            entry.get('auditee', ''),
            entry.get('due_date', ''),
            entry.get('status', ''),
        ], color=row_color, client_key=client_key)

    summary = data.get('summary', {})
    if summary:
        doc.add_paragraph()
        add_section_heading(doc, t("summary", lang), client_key=client_key)
        add_body_text(doc, f"Total Nonconformities: {summary.get('total_nc', 0)}", client_key=client_key)
        add_body_text(doc, f"  Major: {summary.get('major', 0)}", client_key=client_key)
        add_body_text(doc, f"  Minor: {summary.get('minor', 0)}", client_key=client_key)
        add_body_text(doc, f"Opportunities for Improvement: {summary.get('ofi', 0)}", client_key=client_key)
        add_body_text(doc, f"Observations: {summary.get('observations', 0)}", client_key=client_key)
    doc.save(output_path)
    return output_path


def generate_certificate(data, output_path, client_key: str = None):
    """Standalone certificate document — client-aware."""
    client, lang, rtl, header_color = _resolve_client(client_key)
    labels = COVER_LABELS[lang]

    doc = Document()
    setup_document(doc, client_key=client_key)
    add_border_box(doc)
    for _ in range(3):
        doc.add_paragraph()
    if os.path.exists(DEFAULT_LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(DEFAULT_LOGO_PATH, width=Inches(2.0))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(labels["certified"])
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = header_color
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(labels["of_audit"])
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{labels['client']}: {data.get('client_name', 'N/A')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(labels["certified_at"])
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.get('client_name', 'N/A'))
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = header_color
    doc.add_paragraph()
    info_lines = [
        f"{labels['audit_date']}: {data.get('audit_date', 'N/A')}",
        f"{labels['standard']}: {data.get('standard', 'N/A')}",
        f"{labels['scope']}: {data.get('scope', 'N/A')}",
        f"{labels['lead_auditor']}: {data.get('lead_auditor', 'N/A')}",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    decision = data.get('certification_decision', 'Not Certified')
    run = p.add_run(f'{labels["certification_decision"]}: {decision}')
    run.font.size = Pt(14)
    run.bold = True
    if decision == 'Certified':
        run.font.color.rgb = header_color
    elif decision == 'Conditional':
        run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)
    else:
        run.font.color.rgb = TUV_RED
    if data.get('conditions'):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{labels["conditions"]}:')
        run.bold = True
        run.font.size = Pt(10)
        for cond in data['conditions']:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'• {cond}')
            run.font.size = Pt(9)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{labels['issue_date']}: {data.get('issue_date', 'N/A')}     {labels['expiry_date']}: {data.get('expiry_date', 'N/A')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_________________________')
    run.font.color.rgb = header_color
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.get('authorized_signatory', ''))
    run.bold = True
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TÜV AUSTRIA')
    run.font.size = Pt(10)
    run.font.color.rgb = header_color
    doc.save(output_path)
    return output_path


# ── Excel checklist generation (ISO 27001) ─────────────────────────────────

def _generate_checklist_excel(template_path, data, output_dir, safe_name):
    from openpyxl import load_workbook
    from openpyxl.styles import PatternFill
    import shutil

    out_path = os.path.join(output_dir, f'ISO_Checklist_{safe_name}.xlsx')
    shutil.copy(template_path, out_path)

    wb = load_workbook(out_path)
    ws = wb.active

    sections = data.get('sections', [])
    section_map = {}
    for s in sections:
        clause = s.get('clause', '').strip()
        if clause:
            section_map[clause] = s

    filled = 0
    for row in ws.iter_rows(min_row=4, max_row=ws.max_row):
        ref_cell = row[1]
        rate_cell = row[3]
        evidence_cell = row[4]
        notes_cell = row[7] if len(row) > 7 else None

        ref = str(ref_cell.value or '').strip()
        if not ref:
            continue

        match = section_map.get(ref)
        if not match:
            for clause, s in section_map.items():
                if ref.startswith(clause) or clause.startswith(ref):
                    match = s
                    break

        if match:
            status = match.get('status', '')
            rate_map = {
                'Conformant': 'FC',
                'Partially Conformant': 'PF',
                'Non-Conformant': 'NC',
                'Not Reviewed': 'NR',
            }
            rate = rate_map.get(status, '')
            evidence = match.get('evidence', '')
            notes = match.get('notes', '')

            if rate:
                rate_cell.value = rate
            if evidence:
                evidence_cell.value = evidence
            if notes and notes_cell is not None:
                notes_cell.value = notes

            if status == 'Non-Conformant':
                for c in row:
                    c.fill = PatternFill(start_color='FFF0F0', end_color='FFF0F0', fill_type='solid')
            elif status == 'Partially Conformant':
                for c in row:
                    c.fill = PatternFill(start_color='FFF8E1', end_color='FFF8E1', fill_type='solid')
            filled += 1

    wb.save(out_path)
    return out_path


def generate_audit_plan_stage(data, output_path, stage_label, client_key: str = None):
    client = get_client(client_key) if client_key else None
    lang = client.language if client else "en"

    doc = Document()
    setup_document(doc, client_key=client_key)
    stage = data.get('stage', stage_label)
    add_cover_page(doc, f'Audit Plan - {stage}',
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'), data.get('audit_date', 'N/A'),
                   client_key=client_key)

    add_toc(doc)

    add_section_heading(doc, t("audit_objectives", lang), client_key=client_key)
    for obj in data.get('audit_objectives', []):
        add_bullet(doc, obj, client_key=client_key)

    add_section_heading(doc, t("audit_scope", lang), client_key=client_key)
    add_body_text(doc, data.get('audit_scope', 'N/A'), client_key=client_key)

    add_section_heading(doc, t("audit_criteria", lang), client_key=client_key)
    for c in data.get('audit_criteria', []):
        add_bullet(doc, c, client_key=client_key)

    add_section_heading(doc, t("audit_team", lang), client_key=client_key)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, TABLE_HEADERS[lang]["audit_team"], client_key=client_key)
    for member in data.get('audit_team', []):
        add_data_row(table, [member.get('name', ''), member.get('role', ''), str(member.get('days', ''))], client_key=client_key)

    add_section_heading(doc, t("daily_schedule", lang), client_key=client_key)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, TABLE_HEADERS[lang]["daily_schedule"], client_key=client_key)
    for entry in data.get('daily_schedule', []):
        add_data_row(table, [
            str(entry.get('day', '')),
            entry.get('date', ''),
            entry.get('time', ''),
            entry.get('activity', ''),
            entry.get('auditee', ''),
            entry.get('auditor', ''),
            entry.get('clause', ''),
        ], client_key=client_key)

    add_section_heading(doc, t("confidentiality", lang), client_key=client_key)
    add_body_text(doc, data.get('confidentiality', CONFIDENTIALITY[lang]), client_key=client_key)

    add_section_heading(doc, t("language_report", lang), client_key=client_key)
    add_body_text(doc, f"Language: {data.get('language', 'English')}", client_key=client_key)
    add_body_text(doc, f"Report Due: {data.get('report_date', 'N/A')}", client_key=client_key)

    doc.save(output_path)
    return output_path


def generate_audit_plan_stage_1(data, output_path, client_key: str = None):
    return generate_audit_plan_stage(data, output_path, 'Stage 1 - Readiness Review', client_key=client_key)


def generate_audit_plan_stage_2(data, output_path, client_key: str = None):
    return generate_audit_plan_stage(data, output_path, 'Stage 2 - Certification Audit', client_key=client_key)


def generate_participation_list(data, output_path, client_key: str = None):
    client = get_client(client_key) if client_key else None
    lang = client.language if client else "en"

    doc = Document()
    setup_document(doc, client_key=client_key)
    add_cover_page(doc, t("attendance_record", lang),
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'),
                   data.get('audit_date', 'N/A'),
                   client_key=client_key)

    add_section_heading(doc, t("attendance_record", lang), client_key=client_key)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, TABLE_HEADERS[lang]["participants"], client_key=client_key)
    for p in data.get('participants', []):
        add_data_row(table, [
            p.get('name', ''),
            p.get('company', ''),
            p.get('department', ''),
            p.get('closing_meeting', ''),
            p.get('signature', ''),
        ], client_key=client_key)

    notes = data.get('notes', '')
    if notes:
        doc.add_paragraph()
        add_section_heading(doc, t("notes", lang), client_key=client_key)
        add_body_text(doc, notes, client_key=client_key)

    doc.save(output_path)
    return output_path


def generate_certificate_text(data, output_path, client_key: str = None):
    client = get_client(client_key) if client_key else None
    lang = client.language if client else "en"
    labels = COVER_LABELS[lang]

    doc = Document()
    setup_document(doc, client_key=client_key)
    add_cover_page(doc, 'Certificate',
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'),
                   data.get('audit_date', 'N/A'),
                   client_key=client_key)

    fields = [
        (labels["certification_body"], data.get('certificate_number', 'N/A')),
        (labels["client"], data.get('client_name', 'N/A')),
        (labels["standard"], data.get('standard', 'N/A')),
        (labels["audit_date"], data.get('audit_date', 'N/A')),
        (labels["scope"], data.get('scope', 'N/A')),
        (labels["lead_auditor"], data.get('lead_auditor', 'N/A')),
        (labels["certification_body"], data.get('certification_body', 'TÜV AUSTRIA')),
        (labels["certification_decision"], data.get('certification_decision', 'N/A')),
        (labels["issue_date"], data.get('issue_date', 'N/A')),
        (labels["expiry_date"], data.get('expiry_date', 'N/A')),
        ("Authorized Signatory", data.get('authorized_signatory', 'N/A')),
    ]

    add_section_heading(doc, t("certificate_details", lang), client_key=client_key)
    for label, value in fields:
        p = doc.add_paragraph()
        run = p.add_run(f'{label}: ')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = TUV_BLUE
        run = p.add_run(str(value))
        run.font.size = Pt(10)
        p.space_after = Pt(4)

    doc.save(output_path)
    return output_path


def generate_audit_report(data, output_path, client_key: str = None):
    client = get_client(client_key) if client_key else None
    lang = client.language if client else "en"
    standard = data.get('standard', 'the applicable standard')

    doc = Document()
    setup_document(doc, client_key=client_key)
    add_cover_page(doc, 'Audit Report', data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'), data.get('audit_date', 'N/A'),
                   data.get('lead_auditor', ''),
                   client_key=client_key)

    add_toc(doc)

    add_section_heading(doc, t("general_info", lang), client_key=client_key)
    info = [
        ('Report Number', data.get('report_number', 'N/A')),
        (COVER_LABELS[lang]["client"], data.get('client_name', 'N/A')),
        (COVER_LABELS[lang]["audit_date"], data.get('audit_date', 'N/A')),
        (COVER_LABELS[lang]["standard"], data.get('standard', 'N/A')),
        (COVER_LABELS[lang]["scope"], data.get('scope', 'N/A')),
        (COVER_LABELS[lang]["lead_auditor"], data.get('lead_auditor', 'N/A')),
    ]
    for label, val in info:
        add_body_text(doc, f'{label}: {val}', bold=False, client_key=client_key)

    add_section_heading(doc, t("audit_team", lang), client_key=client_key)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_header_row(table, TABLE_HEADERS[lang]["audit_team"], client_key=client_key)
    for member in data.get('audit_team', []):
        add_data_row(table, [member.get('name', ''), member.get('role', ''), str(member.get('days', ''))], client_key=client_key)

    add_section_heading(doc, t("executive_summary", lang), client_key=client_key)
    exec_summary = data.get('findings_summary', '')
    add_body_text(doc, exec_summary, client_key=client_key)
    if data.get('positive_findings'):
        add_sub_heading(doc, t("key_strengths", lang), client_key=client_key)
        for pf in data['positive_findings']:
            add_bullet(doc, pf, client_key=client_key)

    add_section_heading(doc, t("audit_methodology", lang), client_key=client_key)
    default_methodology = {
        'approach': METHODOLOGY[lang]["approach"].format(standard=standard),
        'sampling': METHODOLOGY[lang]["sampling"].format(standard=standard),
        'criteria': METHODOLOGY[lang]["criteria"].format(standard=standard),
        'methods': METHODOLOGY[lang]["methods"].format(standard=standard),
    }
    methodology = data.get('methodology', default_methodology)
    add_sub_heading(doc, t("approach", lang), client_key=client_key)
    add_body_text(doc, methodology.get('approach', ''), client_key=client_key)
    add_sub_heading(doc, t("sampling", lang), client_key=client_key)
    add_body_text(doc, methodology.get('sampling', ''), client_key=client_key)
    add_sub_heading(doc, t("criteria", lang), client_key=client_key)
    add_body_text(doc, methodology.get('criteria', ''), client_key=client_key)
    add_sub_heading(doc, t("methods", lang), client_key=client_key)
    add_body_text(doc, methodology.get('methods', ''), client_key=client_key)

    add_section_heading(doc, t("detailed_findings", lang), client_key=client_key)

    if data.get('positive_findings'):
        add_sub_heading(doc, t("positive_findings", lang), client_key=client_key)
        for pf in data['positive_findings']:
            add_bullet(doc, pf, client_key=client_key)

    if data.get('opportunities_for_improvement'):
        add_sub_heading(doc, t("ofi", lang), client_key=client_key)
        for ofi in data['opportunities_for_improvement']:
            add_bullet(doc, ofi, client_key=client_key)

    if data.get('nonconformities'):
        add_sub_heading(doc, t("nonconformities", lang), client_key=client_key)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        add_header_row(table, TABLE_HEADERS[lang]["nc_table"], client_key=client_key)
        for nc in data['nonconformities']:
            sev = nc.get('severity', '')
            row_color = TUV_RED if sev in ('Major', 'Minor') else None
            add_data_row(table, [
                nc.get('clause', ''),
                sev,
                nc.get('description', ''),
                nc.get('due_date', ''),
            ], color=row_color, client_key=client_key)

    add_section_heading(doc, t("conclusion", lang), client_key=client_key)
    add_body_text(doc, data.get('conclusion', ''), client_key=client_key)

    p = doc.add_paragraph()
    run = p.add_run(f"\nReport Date: {data.get('report_date', 'N/A')}")
    run.font.size = Pt(10)

    doc.save(output_path)
    return output_path


def generate_iso_checklist(data, output_path, template_path=None, client_key: str = None):
    client = get_client(client_key) if client_key else None
    lang = client.language if client else "en"

    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
        setup_document(doc, landscape=True, client_key=client_key)
        sections = data.get('sections', [])

        for table in doc.tables:
            if len(table.rows) < 3:
                continue
            num_cols = len(table.rows[0].cells)

            if num_cols >= 5:
                _inject_into_template_table(table, sections)
                set_col_widths(table, [12, 22, 12, 28, 14, 12], available_inches=9.5)
                break

        overall = data.get('overall_assessment', '')
        if overall:
            for p in doc.paragraphs:
                text = p.text.strip().lower()
                if ('overall' in text or 'assessment' in text) and len(text) < 40:
                    run = p.add_run(f'\n{overall}')
                    run.font.size = Pt(10)
                    break
        doc.save(output_path)
        return output_path

    doc = Document()
    setup_document(doc, landscape=True, client_key=client_key)
    add_cover_page(doc, t("checklist_results", lang),
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'),
                   data.get('audit_date', 'N/A'),
                   data.get('auditor', ''),
                   client_key=client_key)

    add_section_heading(doc, t("checklist_results", lang), client_key=client_key)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    add_header_row(table, TABLE_HEADERS[lang]["checklist"], client_key=client_key)
    for section in data.get('sections', []):
        status = section.get('status', '')
        checkbox = '☐'
        row_color = None
        if status == 'Conformant':
            checkbox = '☒'
        elif status == 'Non-Conformant':
            checkbox = '☒'
            row_color = TUV_RED
        elif status == 'Partially Conformant':
            checkbox = '☒'
            row_color = RGBColor(0xCC, 0x7A, 0x00)
        add_data_row(table, [
            section.get('clause', ''),
            section.get('requirement', ''),
            f'{checkbox} {status}',
            section.get('evidence', ''),
            section.get('notes', ''),
            section.get('reference', ''),
        ], color=row_color, client_key=client_key)

    doc.add_paragraph()
    add_section_heading(doc, t("overall_assessment", lang), client_key=client_key)
    add_body_text(doc, data.get('overall_assessment', ''), client_key=client_key)

    doc.save(output_path)
    return output_path


def _inject_into_template_table(table, sections, data_row_start=1):
    for i, section in enumerate(sections):
        row_idx = data_row_start + i
        if row_idx >= len(table.rows):
            break
        row = table.rows[row_idx]
        status = section.get('status', '')
        checkbox = '☐'
        row_color = None
        if status == 'Conformant':
            checkbox = '☒'
        elif status == 'Non-Conformant':
            checkbox = '☒'
            row_color = TUV_RED
        elif status == 'Partially Conformant':
            checkbox = '☒'
            row_color = RGBColor(0xCC, 0x7A, 0x00)

        col_map = [
            ('clause', 0),
            ('requirement', 1),
            ('title', 1),
            ('status', 2),
            ('evidence', 3),
            ('notes', 4),
            ('reference', 5),
        ]

        cell_values = {}
        for field, col in col_map:
            val = section.get(field, '')
            if field == 'status' and val:
                val = f'{checkbox} {val}'
            if val:
                cell_values[col] = val

        for col_idx, val in cell_values.items():
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(8)
                if row_color and (col_idx == 2 or col_idx == 3):
                    run.font.color.rgb = row_color


# ── Template injection helpers ──────────────────────────────────────────────

def _inject_into_template_by_clause(table, sections, data_start_row, col_clause=0,
                                     col_evidence=None, col_assessment=None,
                                     assessment_map=None, evidence_cols=None):
    """Inject AI evidence/assessment into pre-filled clause rows by matching clause reference."""
    section_map = {}
    for s in sections:
        clause = s.get('clause', '').strip()
        if clause:
            section_map[clause] = s

    filled = 0
    for row_idx in range(data_start_row, len(table.rows)):
        row = table.rows[row_idx]
        if col_clause >= len(row.cells):
            continue
        clause_text = row.cells[col_clause].text.strip()
        if not clause_text:
            continue
        first_line = clause_text.split('\n')[0].strip()
        if not re.search(r'\d+\.\d+', first_line):
            continue

        match = None
        for c, s in section_map.items():
            if first_line.startswith(c) or (len(c) > 3 and first_line.startswith(c)):
                match = s
                break
        if not match:
            for c, s in section_map.items():
                if first_line and c and first_line[0:10] == c[0:10]:
                    match = s
                    break

        if match:
            status = match.get('status', '')
            evidence = match.get('evidence', '')

            ev_cols = evidence_cols if evidence_cols is not None else ([col_evidence] if col_evidence is not None else [])
            for ec in ev_cols:
                if evidence and ec < len(row.cells):
                    cell = row.cells[ec]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    run = p.add_run(evidence)
                    run.font.size = Pt(8)

            if status and col_assessment is not None and col_assessment < len(row.cells):
                if assessment_map:
                    assessment = assessment_map.get(status, '')
                else:
                    assessment = status
                cell = row.cells[col_assessment]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(assessment))
                run.font.size = Pt(8)
                if status == 'Non-Conformant':
                    run.font.color.rgb = TUV_RED
                elif status == 'Partially Conformant':
                    run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)
            filled += 1
    return filled


def _combine_evidence_notes(section):
    """Combine evidence and notes into one cell value for auditor-notes-together templates."""
    evidence = section.get('evidence', '').strip()
    notes = section.get('notes', '').strip()
    if evidence and notes:
        return evidence + '\nNotes: ' + notes
    return evidence or notes


# ── Management Review Minutes (ISO 9.3) ─────────────────────────────────────

def generate_management_review_minutes(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)

    doc = Document()
    setup_document(doc, client_key=client_key)
    add_cover_page(doc, 'Management Review Minutes',
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'),
                   data.get('review_date', 'N/A'),
                   client_key=client_key)

    add_section_heading(doc, 'Review Information', client_key=client_key)
    info = [
        ('Chairperson', data.get('chairperson', 'N/A')),
        ('Review Date', data.get('review_date', 'N/A')),
        ('Next Review Date', data.get('next_review_date', 'N/A')),
    ]
    for label, val in info:
        add_body_text(doc, f'{label}: {val}', client_key=client_key)

    add_section_heading(doc, 'Attendees', client_key=client_key)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_header_row(table, ['Name', 'Role', 'Department'], client_key=client_key)
    for att in data.get('attendees', []):
        add_data_row(table, [att.get('name', ''), att.get('role', ''), att.get('department', '')], client_key=client_key)

    add_section_heading(doc, 'Agenda and Discussion', client_key=client_key)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_header_row(table, ['Agenda Item', 'Presented By', 'Discussion Summary'], client_key=client_key)
    for item in data.get('agenda_items', []):
        add_data_row(table, [item.get('item', ''), item.get('presented_by', ''), item.get('discussion', '')], client_key=client_key)

    decisions = data.get('decisions', [])
    if decisions:
        add_section_heading(doc, 'Decisions', client_key=client_key)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        add_header_row(table, ['Decision', 'Rationale', 'Owner'], client_key=client_key)
        for dec in decisions:
            add_data_row(table, [dec.get('decision', ''), dec.get('rationale', ''), dec.get('owner', '')], client_key=client_key)

    actions = data.get('action_items', [])
    if actions:
        add_section_heading(doc, 'Action Items', client_key=client_key)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        add_header_row(table, ['Action', 'Owner', 'Due Date', 'Status'], client_key=client_key)
        for act in actions:
            add_data_row(table, [act.get('action', ''), act.get('owner', ''), act.get('due_date', ''), act.get('status', '')], client_key=client_key)

    add_section_heading(doc, 'Review Inputs', client_key=client_key)
    add_body_text(doc, data.get('review_inputs', ''), client_key=client_key)

    add_section_heading(doc, 'Review Outputs', client_key=client_key)
    add_body_text(doc, data.get('review_outputs', ''), client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Corrective Action Report (ISO 10.1) ─────────────────────────────────────

def generate_corrective_action_report(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)

    doc = Document()
    setup_document(doc, client_key=client_key)
    add_cover_page(doc, 'Corrective Action Report',
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'),
                   data.get('audit_date', 'N/A'),
                   client_key=client_key)

    add_section_heading(doc, 'Nonconformity Information', client_key=client_key)
    fields = [
        ('CAR Number', data.get('car_number', 'N/A')),
        ('NC Reference', data.get('nc_reference', 'N/A')),
        ('ISO Clause', data.get('clause', 'N/A')),
        ('Severity', data.get('severity', 'N/A')),
        ('Status', data.get('status', 'Open')),
    ]
    for label, val in fields:
        add_body_text(doc, f'{label}: {val}', client_key=client_key)

    add_section_heading(doc, 'Problem Description', client_key=client_key)
    add_body_text(doc, data.get('problem_description', ''), client_key=client_key)

    add_section_heading(doc, 'Root Cause Analysis', client_key=client_key)
    add_body_text(doc, data.get('root_cause', ''), client_key=client_key)

    ca = data.get('containment_actions', [])
    if ca:
        add_section_heading(doc, 'Containment Actions', client_key=client_key)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        add_header_row(table, ['Action', 'Owner', 'Due Date'], client_key=client_key)
        for act in ca:
            add_data_row(table, [act.get('action', ''), act.get('owner', ''), act.get('due_date', '')], client_key=client_key)

    corr = data.get('corrective_actions', [])
    if corr:
        add_section_heading(doc, 'Corrective Actions', client_key=client_key)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        add_header_row(table, ['Action', 'Owner', 'Due Date'], client_key=client_key)
        for act in corr:
            add_data_row(table, [act.get('action', ''), act.get('owner', ''), act.get('due_date', '')], client_key=client_key)

    prev = data.get('preventive_actions', [])
    if prev:
        add_section_heading(doc, 'Preventive Actions', client_key=client_key)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        add_header_row(table, ['Action', 'Owner', 'Due Date'], client_key=client_key)
        for act in prev:
            add_data_row(table, [act.get('action', ''), act.get('owner', ''), act.get('due_date', '')], client_key=client_key)

    add_section_heading(doc, 'Verification', client_key=client_key)
    add_body_text(doc, f'Method: {data.get("verification_method", "")}', client_key=client_key)
    add_body_text(doc, f'Reviewed By: {data.get("reviewed_by", "")}', client_key=client_key)
    add_body_text(doc, f'Target Closure: {data.get("closure_date", "")}', client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Gap Analysis Report (Pre-Audit) ─────────────────────────────────────────

def generate_gap_analysis_report(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)

    doc = Document()
    setup_document(doc, landscape=True, client_key=client_key)
    add_cover_page(doc, 'Gap Analysis Report',
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'),
                   data.get('assessment_date', 'N/A'),
                   client_key=client_key)

    add_section_heading(doc, 'Assessment Methodology', client_key=client_key)
    add_body_text(doc, data.get('methodology', ''), client_key=client_key)

    add_section_heading(doc, 'Clause-by-Clause Gap Assessment', client_key=client_key)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    add_header_row(table, ['Clause', 'Title', 'Requirement', 'Status', 'Gap Description', 'Recommended Action', 'Priority'], client_key=client_key)
    set_col_widths(table, [8, 12, 20, 12, 20, 20, 8], available_inches=9.5)

    for section in data.get('sections', []):
        status = section.get('status', '')
        row_color = None
        if status == 'Non-Conformant':
            row_color = TUV_RED
        elif status == 'Partially Conformant':
            row_color = RGBColor(0xCC, 0x7A, 0x00)
        elif status == 'Conformant':
            row_color = RGBColor(0x00, 0x80, 0x00)
        add_data_row(table, [
            section.get('clause', ''),
            section.get('title', ''),
            section.get('requirement', ''),
            status,
            section.get('gap_description', ''),
            section.get('recommended_action', ''),
            section.get('priority', ''),
        ], color=row_color, client_key=client_key)

    summary = data.get('summary', {})
    if summary:
        doc.add_paragraph()
        add_section_heading(doc, 'Summary', client_key=client_key)
        add_body_text(doc, f"Total Clauses Assessed: {summary.get('total_clauses', 0)}", client_key=client_key)
        add_body_text(doc, f"Conformant: {summary.get('conformant', 0)}", client_key=client_key)
        add_body_text(doc, f"Partially Conformant: {summary.get('partially_conformant', 0)}", client_key=client_key)
        add_body_text(doc, f"Non-Conformant: {summary.get('non_conformant', 0)}", client_key=client_key)
        add_body_text(doc, f"Not Reviewed: {summary.get('not_reviewed', 0)}", client_key=client_key)
        add_body_text(doc, f"Overall Readiness: {summary.get('overall_readiness', 'N/A')}", client_key=client_key)

    overall = data.get('overall_assessment', '')
    if overall:
        add_section_heading(doc, 'Overall Assessment', client_key=client_key)
        add_body_text(doc, overall, client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Statement of Applicability (Annex A — ISO 27001 / 42001) ────────────────

def generate_statement_of_applicability(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)

    doc = Document()
    setup_document(doc, landscape=True, client_key=client_key)
    add_cover_page(doc, 'Statement of Applicability (SoA)',
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'),
                   data.get('date', 'N/A'),
                   client_key=client_key)

    add_section_heading(doc, 'Risk Assessment Reference', client_key=client_key)
    add_body_text(doc, data.get('based_on_risk_assessment', ''), client_key=client_key)

    controls = data.get('controls', [])
    if controls:
        add_section_heading(doc, 'Annex A Controls', client_key=client_key)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        add_header_row(table, ['Control Ref', 'Control Title', 'Applicability', 'Justification', 'Implementation Status', 'Responsible'], client_key=client_key)
        set_col_widths(table, [8, 15, 10, 35, 15, 12], available_inches=9.5)

        for ctrl in controls:
            status = ctrl.get('implementation_status', '')
            row_color = None
            if status == 'Not Implemented':
                row_color = TUV_RED
            elif status == 'Planned':
                row_color = RGBColor(0xCC, 0x7A, 0x00)
            elif status == 'Implemented':
                row_color = RGBColor(0x00, 0x80, 0x00)

            applic = ctrl.get('applicability', '')
            if applic == 'Not Applicable':
                row_color = RGBColor(0x66, 0x66, 0x66)

            add_data_row(table, [
                ctrl.get('control_ref', ''),
                ctrl.get('control_title', ''),
                applic,
                ctrl.get('justification', ''),
                status,
                ctrl.get('responsible', ''),
            ], color=row_color, client_key=client_key)

    summary = data.get('summary', {})
    if summary:
        doc.add_paragraph()
        add_section_heading(doc, 'Summary', client_key=client_key)
        add_body_text(doc, f"Total Controls: {summary.get('total_controls', 0)}", client_key=client_key)
        add_body_text(doc, f"Applicable: {summary.get('applicable', 0)}", client_key=client_key)
        add_body_text(doc, f"Not Applicable: {summary.get('not_applicable', 0)}", client_key=client_key)
        add_body_text(doc, f"Implemented: {summary.get('implemented', 0)}", client_key=client_key)
        add_body_text(doc, f"Not Implemented: {summary.get('not_implemented', 0)}", client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Business Impact Analysis (ISO 22301 Clause 8.2.1) ───────────────────────

def generate_business_impact_analysis(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)

    doc = Document()
    setup_document(doc, client_key=client_key)
    add_cover_page(doc, 'Business Impact Analysis (BIA)',
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'),
                   data.get('assessment_date', 'N/A'),
                   client_key=client_key)

    add_section_heading(doc, 'Methodology', client_key=client_key)
    add_body_text(doc, data.get('methodology', ''), client_key=client_key)

    activities = data.get('critical_activities', [])
    if activities:
        add_section_heading(doc, 'Critical Activities Assessment', client_key=client_key)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        add_header_row(table, ['Activity', 'RTO', 'RPO', 'MTD', 'Impact Criteria', 'Dependencies', 'Recovery Strategy', 'Priority'], client_key=client_key)

        for act in activities:
            priority = act.get('priority', '')
            row_color = None
            if priority == 'Critical':
                row_color = RGBColor(0xCC, 0x00, 0x00)
            elif priority == 'High':
                row_color = RGBColor(0xCC, 0x7A, 0x00)
            elif priority == 'Medium':
                row_color = RGBColor(0x00, 0x80, 0x00)

            add_data_row(table, [
                act.get('activity', ''),
                act.get('rto', ''),
                act.get('rpo', ''),
                act.get('mtd', ''),
                act.get('impact_criteria', ''),
                act.get('dependencies', ''),
                act.get('recovery_strategy', ''),
                priority,
            ], color=row_color, client_key=client_key)

    summary = data.get('summary', {})
    if summary:
        doc.add_paragraph()
        add_section_heading(doc, 'Summary', client_key=client_key)
        add_body_text(doc, f"Total Activities Assessed: {summary.get('total_activities', 0)}", client_key=client_key)
        add_body_text(doc, f"Critical: {summary.get('critical', 0)}", client_key=client_key)
        add_body_text(doc, f"High: {summary.get('high', 0)}", client_key=client_key)
        add_body_text(doc, f"Medium: {summary.get('medium', 0)}", client_key=client_key)
        add_body_text(doc, f"Low: {summary.get('low', 0)}", client_key=client_key)

    findings = data.get('overall_findings', '')
    if findings:
        add_section_heading(doc, 'Overall Findings', client_key=client_key)
        add_body_text(doc, findings, client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Records of Processing Activities (ISO 27701 Clause 7.6) ────────────────

def generate_records_of_processing_activities(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)

    doc = Document()
    setup_document(doc, landscape=True, client_key=client_key)
    add_cover_page(doc, 'Records of Processing Activities (ROPA)',
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'),
                   data.get('date', 'N/A'),
                   client_key=client_key)

    add_section_heading(doc, 'Data Controller Information', client_key=client_key)
    add_body_text(doc, f"Data Controller: {data.get('data_controller', 'N/A')}", client_key=client_key)
    add_body_text(doc, f"DPO: {data.get('data_protection_officer', 'N/A')}", client_key=client_key)

    activities = data.get('processing_activities', [])
    if activities:
        add_section_heading(doc, 'Processing Activities Register', client_key=client_key)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        headers = ['Activity ID', 'Activity Name', 'Purpose', 'Data Subjects', 'Data Categories', 'Retention', 'Cross-Border Transfer', 'Security Measures']
        add_header_row(table, headers, client_key=client_key)

        for act in activities:
            add_data_row(table, [
                act.get('activity_id', ''),
                act.get('activity_name', ''),
                act.get('purpose', ''),
                act.get('data_subjects', ''),
                act.get('personal_data_categories', ''),
                act.get('retention_period', ''),
                act.get('cross_border_transfer', ''),
                act.get('security_measures', ''),
            ], client_key=client_key)

    summary = data.get('summary', {})
    if summary:
        doc.add_paragraph()
        add_section_heading(doc, 'Summary', client_key=client_key)
        add_body_text(doc, f"Total Processing Activities: {summary.get('total_activities', 0)}", client_key=client_key)
        add_body_text(doc, f"Cross-Border Transfers: {summary.get('has_cross_border_transfers', 'N/A')}", client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Risk Treatment Plan (ISO 27001 Clause 8.3) ────────────────────────────

def generate_risk_treatment_plan(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)

    doc = Document()
    setup_document(doc, landscape=True, client_key=client_key)
    add_cover_page(doc, 'Risk Treatment Plan',
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'),
                   data.get('date', 'N/A'),
                   client_key=client_key)

    add_section_heading(doc, 'Risk Assessment Reference', client_key=client_key)
    add_body_text(doc, data.get('risk_assessment_reference', ''), client_key=client_key)

    risks = data.get('risks', [])
    if risks:
        add_section_heading(doc, 'Risk Treatment Register', client_key=client_key)
        table = doc.add_table(rows=1, cols=11)
        table.style = 'Table Grid'
        add_header_row(table, ['Risk ID', 'Description', 'Source', 'Likelihood', 'Impact', 'Risk Level',
                               'Treatment Option', 'Treatment Details', 'Selected Controls', 'Risk Owner', 'Target Date', 'Status'], client_key=client_key)

        for risk in risks:
            level = risk.get('risk_level', '')
            row_color = None
            if level == 'Critical':
                row_color = TUV_RED
            elif level == 'High':
                row_color = RGBColor(0xCC, 0x7A, 0x00)
            elif level == 'Medium':
                row_color = RGBColor(0x00, 0x80, 0x00)

            add_data_row(table, [
                risk.get('risk_id', ''),
                risk.get('risk_description', ''),
                risk.get('source', ''),
                risk.get('likelihood', ''),
                risk.get('impact', ''),
                level,
                risk.get('treatment_option', ''),
                risk.get('treatment_details', ''),
                risk.get('selected_controls', ''),
                risk.get('risk_owner', ''),
                risk.get('target_date', ''),
                risk.get('status', ''),
            ], color=row_color, client_key=client_key)

    summary = data.get('summary', {})
    if summary:
        doc.add_paragraph()
        add_section_heading(doc, 'Summary', client_key=client_key)
        add_body_text(doc, f"Total Risks: {summary.get('total_risks', 0)}", client_key=client_key)
        add_body_text(doc, f"Critical: {summary.get('critical', 0)}", client_key=client_key)
        add_body_text(doc, f"High: {summary.get('high', 0)}", client_key=client_key)
        add_body_text(doc, f"Medium: {summary.get('medium', 0)}", client_key=client_key)
        add_body_text(doc, f"Low: {summary.get('low', 0)}", client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Incident Investigation Report (ISO 45001 Clause 10.2) ─────────────────

def generate_incident_investigation_report(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)

    doc = Document()
    setup_document(doc, client_key=client_key)
    add_cover_page(doc, 'Incident Investigation Report',
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'),
                   data.get('report_date', 'N/A'),
                   client_key=client_key)

    add_section_heading(doc, 'Incident Information', client_key=client_key)
    fields = [
        ('Incident Date', data.get('incident_date', 'N/A')),
        ('Report Date', data.get('report_date', 'N/A')),
        ('Location', data.get('location', 'N/A')),
        ('Incident Type', data.get('incident_type', 'N/A')),
        ('Severity', data.get('severity', 'N/A')),
        ('Status', data.get('status', 'Open')),
    ]
    for label, val in fields:
        add_body_text(doc, f'{label}: {val}', client_key=client_key)

    desc = data.get('incident_description', '')
    if desc:
        add_section_heading(doc, 'Incident Description', client_key=client_key)
        add_body_text(doc, desc, client_key=client_key)

    team = data.get('investigation_team', [])
    if team:
        add_section_heading(doc, 'Investigation Team', client_key=client_key)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        add_header_row(table, ['Name', 'Role'], client_key=client_key)
        for m in team:
            add_data_row(table, [m.get('name', ''), m.get('role', '')], client_key=client_key)

    rc = data.get('root_cause', '')
    if rc:
        add_section_heading(doc, 'Root Cause Analysis', client_key=client_key)
        add_body_text(doc, rc, client_key=client_key)

    imm = data.get('immediate_actions', [])
    if imm:
        add_section_heading(doc, 'Immediate Actions Taken', client_key=client_key)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        add_header_row(table, ['Action', 'Owner', 'Due Date'], client_key=client_key)
        for act in imm:
            add_data_row(table, [act.get('action', ''), act.get('owner', ''), act.get('due_date', '')], client_key=client_key)

    corr = data.get('corrective_actions', [])
    if corr:
        add_section_heading(doc, 'Corrective Actions', client_key=client_key)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        add_header_row(table, ['Action', 'Owner', 'Due Date'], client_key=client_key)
        for act in corr:
            add_data_row(table, [act.get('action', ''), act.get('owner', ''), act.get('due_date', '')], client_key=client_key)

    lessons = data.get('lessons_learned', [])
    if lessons:
        add_section_heading(doc, 'Lessons Learned', client_key=client_key)
        for i, lesson in enumerate(lessons, 1):
            add_body_text(doc, f'{i}. {lesson}', client_key=client_key)

    recs = data.get('recommendations', [])
    if recs:
        add_section_heading(doc, 'Recommendations', client_key=client_key)
        for i, rec in enumerate(recs, 1):
            add_body_text(doc, f'{i}. {rec}', client_key=client_key)

    add_section_heading(doc, 'Review', client_key=client_key)
    add_body_text(doc, f'Reviewed By: {data.get("reviewed_by", "N/A")}', client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Internal Audit Program (ISO Clause 9.2) ────────────────────────────────

def generate_internal_audit_program(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)

    doc = Document()
    setup_document(doc, client_key=client_key)
    add_cover_page(doc, 'Internal Audit Program',
                   data.get('client_name', 'N/A'),
                   data.get('standard', 'N/A'),
                   data.get('program_year', 'N/A'),
                   client_key=client_key)

    add_section_heading(doc, 'Program Information', client_key=client_key)
    add_body_text(doc, f'Audit Manager: {data.get("audit_manager", "N/A")}', client_key=client_key)
    add_body_text(doc, f'Program Year: {data.get("program_year", "N/A")}', client_key=client_key)

    audits = data.get('audits', [])
    if audits:
        add_section_heading(doc, 'Audit Schedule', client_key=client_key)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        add_header_row(table, ['Audit ID', 'Scope', 'Type', 'Planned Date', 'Auditor', 'Department', 'Status', 'Findings'], client_key=client_key)

        for a in audits:
            status = a.get('status', '')
            row_color = None
            if status == 'Completed':
                row_color = RGBColor(0x00, 0x80, 0x00)
            elif status == 'In Progress':
                row_color = RGBColor(0xCC, 0x7A, 0x00)
            elif status == 'Cancelled':
                row_color = RGBColor(0x99, 0x99, 0x99)

            add_data_row(table, [
                a.get('audit_id', ''),
                a.get('scope', ''),
                a.get('audit_type', ''),
                a.get('planned_date', ''),
                a.get('auditor', ''),
                a.get('auditee_department', ''),
                status,
                str(a.get('findings_count', 0)),
            ], color=row_color, client_key=client_key)

    summary = data.get('summary', {})
    if summary:
        doc.add_paragraph()
        add_section_heading(doc, 'Summary', client_key=client_key)
        add_body_text(doc, f"Total Audits: {summary.get('total_audits', 0)}", client_key=client_key)
        add_body_text(doc, f"Planned: {summary.get('planned', 0)}", client_key=client_key)
        add_body_text(doc, f"In Progress: {summary.get('in_progress', 0)}", client_key=client_key)
        add_body_text(doc, f"Completed: {summary.get('completed', 0)}", client_key=client_key)
        add_body_text(doc, f"Cancelled: {summary.get('cancelled', 0)}", client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Environmental Aspect Register (ISO 14001 Clause 6.1.2) ─────────────────

def generate_environmental_aspect_register(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)
    doc = Document()
    setup_document(doc, landscape=True, client_key=client_key)
    add_cover_page(doc, 'Environmental Aspect Register',
                   data.get('client_name', 'N/A'), data.get('standard', 'N/A'),
                   data.get('date', 'N/A'), client_key=client_key)

    add_section_heading(doc, 'Methodology', client_key=client_key)
    add_body_text(doc, data.get('methodology', ''), client_key=client_key)

    aspects = data.get('aspects', [])
    if aspects:
        add_section_heading(doc, 'Environmental Aspects Register', client_key=client_key)
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        add_header_row(table, ['Aspect ID', 'Activity', 'Aspect', 'Environmental Impact', 'Type',
                               'Significance', 'Control Measures', 'Legal Requirement', 'Evaluation'], client_key=client_key)
        for a in aspects:
            sig = a.get('significance', '')
            rc = TUV_RED if sig == 'Critical' else (RGBColor(0xCC, 0x7A, 0x00) if sig == 'High' else None)
            add_data_row(table, [
                a.get('aspect_id', ''), a.get('activity', ''), a.get('aspect', ''),
                a.get('environmental_impact', ''), a.get('impact_type', ''),
                sig, a.get('control_measures', ''), a.get('legal_requirement', ''), a.get('evaluation', ''),
            ], color=rc, client_key=client_key)

    summary = data.get('summary', {})
    if summary:
        doc.add_paragraph()
        add_section_heading(doc, 'Summary', client_key=client_key)
        for k, v in [('Total Aspects', 'total_aspects'), ('Critical', 'critical'),
                      ('High', 'high'), ('Medium', 'medium'), ('Low', 'low')]:
            add_body_text(doc, f'{k}: {summary.get(v, 0)}', client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Hazard Identification Register (ISO 45001 Clause 6.1.2) ───────────────

def generate_hazard_identification_register(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)
    doc = Document()
    setup_document(doc, landscape=True, client_key=client_key)
    add_cover_page(doc, 'Hazard Identification Register',
                   data.get('client_name', 'N/A'), data.get('standard', 'N/A'),
                   data.get('date', 'N/A'), client_key=client_key)

    add_section_heading(doc, 'Methodology', client_key=client_key)
    add_body_text(doc, data.get('methodology', ''), client_key=client_key)

    hazards = data.get('hazards', [])
    if hazards:
        add_section_heading(doc, 'Hazard Register', client_key=client_key)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        add_header_row(table, ['Hazard ID', 'Activity', 'Hazard', 'Associated Risk', 'Existing Controls',
                               'Risk Level', 'Additional Controls', 'Hierarchy of Control'], client_key=client_key)
        for h in hazards:
            lv = h.get('risk_level', '')
            rc = TUV_RED if lv == 'Critical' else (RGBColor(0xCC, 0x7A, 0x00) if lv == 'High' else None)
            add_data_row(table, [
                h.get('hazard_id', ''), h.get('activity', ''), h.get('hazard', ''),
                h.get('associated_risk', ''), h.get('existing_controls', ''),
                lv, h.get('additional_controls', ''), h.get('hierarchy_of_control', ''),
            ], color=rc, client_key=client_key)

    summary = data.get('summary', {})
    if summary:
        doc.add_paragraph()
        add_section_heading(doc, 'Summary', client_key=client_key)
        for k, v in [('Total Hazards', 'total_hazards'), ('Critical', 'critical'),
                      ('High', 'high'), ('Medium', 'medium'), ('Low', 'low')]:
            add_body_text(doc, f'{k}: {summary.get(v, 0)}', client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Energy Review + EnB + EnPI (ISO 50001 Clause 6.3) ──────────────────────

def generate_energy_review(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)
    doc = Document()
    setup_document(doc, client_key=client_key)
    add_cover_page(doc, 'Energy Review',
                   data.get('client_name', 'N/A'), data.get('standard', 'N/A'),
                   data.get('date', 'N/A'), client_key=client_key)

    add_section_heading(doc, 'Review Period', client_key=client_key)
    add_body_text(doc, data.get('review_period', 'N/A'), client_key=client_key)
    add_section_heading(doc, 'Methodology', client_key=client_key)
    add_body_text(doc, data.get('methodology', ''), client_key=client_key)

    sources = data.get('energy_sources', [])
    if sources:
        add_section_heading(doc, 'Energy Sources & Consumption', client_key=client_key)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        add_header_row(table, ['Energy Source', 'Annual Consumption', 'Annual Cost', 'Trend', 'Notes'], client_key=client_key)
        for s in sources:
            add_data_row(table, [s.get('source', ''), s.get('consumption', ''), s.get('cost', ''),
                                 s.get('trend', ''), s.get('notes', '')], client_key=client_key)

    seus = data.get('significant_uses', [])
    if seus:
        doc.add_paragraph()
        add_section_heading(doc, 'Significant Energy Uses (SEUs) with EnPI & EnB', client_key=client_key)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        add_header_row(table, ['SEU ID', 'Equipment / Process', 'Energy Source', 'Consumption',
                               'Relevant Variables', 'EnPI', 'Energy Baseline (EnB)', 'Current Performance'], client_key=client_key)
        for seu in seus:
            add_data_row(table, [
                seu.get('use_id', ''), seu.get('equipment', ''), seu.get('energy_source', ''),
                seu.get('consumption', ''), seu.get('variables', ''), seu.get('enpi', ''),
                seu.get('baseline', ''), seu.get('current_performance', ''),
            ], client_key=client_key)

    summary = data.get('summary', {})
    if summary:
        doc.add_paragraph()
        add_section_heading(doc, 'Summary', client_key=client_key)
        add_body_text(doc, f"Energy Sources: {summary.get('total_energy_sources', 0)}", client_key=client_key)
        add_body_text(doc, f"Significant Energy Uses: {summary.get('total_seus', 0)}", client_key=client_key)
        add_body_text(doc, f"Total Annual Energy Cost: {summary.get('total_energy_cost', 'N/A')}", client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Compliance Obligations Register (ISO 37301 / 14001 Clause 6.1.3) ───────

def generate_compliance_obligations_register(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)
    doc = Document()
    setup_document(doc, landscape=True, client_key=client_key)
    add_cover_page(doc, 'Compliance Obligations Register',
                   data.get('client_name', 'N/A'), data.get('standard', 'N/A'),
                   data.get('date', 'N/A'), client_key=client_key)

    add_section_heading(doc, 'Methodology', client_key=client_key)
    add_body_text(doc, data.get('methodology', ''), client_key=client_key)

    obligations = data.get('obligations', [])
    if obligations:
        add_section_heading(doc, 'Compliance Obligations', client_key=client_key)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        add_header_row(table, ['Obligation ID', 'Type', 'Source', 'Requirement', 'Applicability',
                               'Compliance Status', 'Evidence', 'Due Date'], client_key=client_key)
        for o in obligations:
            st = o.get('compliance_status', '')
            rc = TUV_RED if st == 'Non-Compliant' else (RGBColor(0xCC, 0x7A, 0x00) if st == 'Partially Compliant' else None)
            add_data_row(table, [
                o.get('obligation_id', ''), o.get('obligation_type', ''), o.get('source', ''),
                o.get('requirement', ''), o.get('applicability', ''),
                st, o.get('evidence', ''), o.get('due_date', ''),
            ], color=rc, client_key=client_key)

    summary = data.get('summary', {})
    if summary:
        doc.add_paragraph()
        add_section_heading(doc, 'Summary', client_key=client_key)
        for k, v in [('Total Obligations', 'total_obligations'), ('Compliant', 'compliant'),
                      ('Partially Compliant', 'partially_compliant'), ('Non-Compliant', 'non_compliant'),
                      ('Not Assessed', 'not_assessed')]:
            add_body_text(doc, f'{k}: {summary.get(v, 0)}', client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Service Portfolio & SLAs (ISO 20000-1 Clause 7.2) ─────────────────────

def generate_service_portfolio(data, output_path, client_key: str = None):
    client, lang, rtl, header_color = _resolve_client(client_key)
    doc = Document()
    setup_document(doc, landscape=True, client_key=client_key)
    add_cover_page(doc, 'Service Portfolio & SLAs',
                   data.get('client_name', 'N/A'), data.get('standard', 'N/A'),
                   data.get('date', 'N/A'), client_key=client_key)

    add_section_heading(doc, 'Portfolio Manager', client_key=client_key)
    add_body_text(doc, data.get('portfolio_manager', 'N/A'), client_key=client_key)

    services = data.get('services', [])
    if services:
        add_section_heading(doc, 'Service Portfolio & SLA Register', client_key=client_key)
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        add_header_row(table, ['Service ID', 'Service Name', 'Description', 'Category', 'Status',
                               'SLA — Uptime', 'SLA — Response Time', 'SLA — Resolution Time', 'Service Owner'], client_key=client_key)
        for s in services:
            st = s.get('status', '')
            rc = RGBColor(0x00, 0x80, 0x00) if st == 'Active' else (RGBColor(0xCC, 0x7A, 0x00) if st == 'In Development' else None)
            add_data_row(table, [
                s.get('service_id', ''), s.get('service_name', ''), s.get('description', ''),
                s.get('category', ''), st,
                s.get('sla_uptime', ''), s.get('sla_response_time', ''), s.get('sla_resolution_time', ''),
                s.get('service_owner', ''),
            ], color=rc, client_key=client_key)

    summary = data.get('summary', {})
    if summary:
        doc.add_paragraph()
        add_section_heading(doc, 'Summary', client_key=client_key)
        for k, v in [('Total Services', 'total_services'), ('Active', 'active'),
                      ('In Development', 'in_development'), ('Retired', 'retired'), ('Planned', 'planned')]:
            add_body_text(doc, f'{k}: {summary.get(v, 0)}', client_key=client_key)

    doc.save(output_path)
    return output_path


# ── Generator registry ─────────────────────────────────────────────────────

GENERATORS = {
    'Audit_Plan_Stage_1': generate_audit_plan_stage_1,
    'Audit_Plan_Stage_2': generate_audit_plan_stage_2,
    'Participation_List': generate_participation_list,
    'Audit_Report': generate_audit_report,
    'ISO_Checklist': generate_iso_checklist,
    'Certificate_Text': generate_certificate_text,
    'TNL': generate_tnl,
    'Certificate': generate_certificate,
    'Management_Review_Minutes': generate_management_review_minutes,
    'Corrective_Action_Report': generate_corrective_action_report,
    'Gap_Analysis_Report': generate_gap_analysis_report,
    'Statement_of_Applicability': generate_statement_of_applicability,
    'Business_Impact_Analysis': generate_business_impact_analysis,
    'Records_of_Processing_Activities': generate_records_of_processing_activities,
    'Risk_Treatment_Plan': generate_risk_treatment_plan,
    'Incident_Investigation_Report': generate_incident_investigation_report,
    'Internal_Audit_Program': generate_internal_audit_program,
    'Environmental_Aspect_Register': generate_environmental_aspect_register,
    'Hazard_Identification_Register': generate_hazard_identification_register,
    'Energy_Review': generate_energy_review,
    'Compliance_Obligations_Register': generate_compliance_obligations_register,
    'Service_Portfolio': generate_service_portfolio,
}


def generate_document_file(doc_type, data, output_dir, template_path=None, standard_key=None, client_key: str = None):
    from app.services.template_manager import get_template_path, get_checklist_is_excel

    os.makedirs(output_dir, exist_ok=True)
    safe_name = sanitize_filename(data.get('client_name', 'Client'))[:60]

    effective_template = template_path if (template_path and os.path.exists(template_path)) else None
    if not effective_template and doc_type == 'ISO_Checklist':
        effective_template = get_template_path(doc_type, standard_key)

    if doc_type == 'ISO_Checklist' and effective_template:
        if get_checklist_is_excel(standard_key):
            result = _generate_checklist_excel(effective_template, data, output_dir, safe_name)
            if result and os.path.exists(result):
                return result
            raise RuntimeError(f"Excel checklist generation failed for {doc_type}")
        filename = f'{doc_type}_{safe_name}.docx'
        output_path = os.path.join(output_dir, filename)
        generator = GENERATORS.get(doc_type)
        if generator:
            result = generator(data, output_path, effective_template, client_key=client_key)
            if result and os.path.exists(result):
                return result
        raise RuntimeError(f"Document generation failed for {doc_type}: no generator produced output")

    filename = f'{doc_type}_{safe_name}.docx'
    output_path = os.path.join(output_dir, filename)
    generator = GENERATORS.get(doc_type)
    if generator:
        result = generator(data, output_path, client_key=client_key)
        if result and os.path.exists(result):
            return result
        raise RuntimeError(f"Document generation failed for {doc_type}: generator produced no output")
    raise RuntimeError(f"Document generation failed: no generator registered for doc_type '{doc_type}'")
