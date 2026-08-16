"""Generic TÜV template filler — fills fixed forms without altering layout.

Loads a TÜV .docx/.xlsx form and injects AI-generated data into labeled fields,
clause tables, and findings/NC tables. Layout is preserved 100%.
"""

import logging

from docx import Document
from docx.shared import Pt, RGBColor

from app.services.template_manager import find_table_by_header

logger = logging.getLogger(__name__)

# Label → (table_idx_search, key) candidates for the company/audit info block
FIELD_LABELS = {
    "company": ["Company name", "Organization name", "1. Company information"],
    "address": ["Address", "Full address"],
    "representative": ["Company representative", "Organization representative"],
    "email": ["E-mail", "email"],
    "scope": ["Scope of certification", "Certification Scope"],
    "audit_date": ["Audit date", "Audit Date"],
    "standard": ["Standard 1", "Covered Standards", "Audit standard"],
    "lead_auditor": ["Lead Auditor", "LA"],
    "audit_type": ["Audit type", "Audit Type"],
}


def _set_cell_if_label(cell, label_key, value, data):
    """If a cell's text matches a known field label, append the value on next line."""
    text = cell.text.strip()
    low = text.lower()
    for key, labels in FIELD_LABELS.items():
        if key == label_key:
            continue
        for lbl in labels:
            if low.startswith(lbl.lower()) and value is not None:
                # Append value after label
                cell.text = text
                p = cell.add_paragraph()
                run = p.add_run(str(value))
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                return True
    return False


def fill_info_block(doc, data: dict):
    """Fill company/audit info labels across all tables and paragraphs."""
    labels = {
        "Company name": data.get("client_name"),
        "Company representative": data.get("client_representative"),
        "E-mail": data.get("client_email"),
        "Scope of certification": data.get("scope"),
        "Lead Auditor": data.get("lead_auditor"),
        "Standard 1": data.get("standard"),
        "Audit date": data.get("audit_date"),
        "Audit type": data.get("audit_type"),
    }
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                for lbl, val in labels.items():
                    if val and text.startswith(lbl):
                        p = cell.add_paragraph()
                        run = p.add_run(str(val))
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                        break
    # Paragraph fields (single-line labels like "Company name")
    for p in doc.paragraphs:
        text = p.text.strip()
        for lbl, val in labels.items():
            if val and text.startswith(lbl + ":") and len(text) < len(lbl) + 20:
                run = p.add_run(' ' + str(val))
                run.font.size = Pt(9)
                break


def fill_clause_ratings(doc, sections: list[dict], rating_map: dict | None = None):
    """Fill clause rating + evidence by matching clause refs in checklist tables."""
    if not sections:
        return 0
    rating_map = rating_map or {
        "Conformant": 1,
        "Partially Conformant": 2,
        "Non-Conformant": 3,
    }
    section_map = {}
    for s in sections:
        c = str(s.get("clause", "")).strip()
        if c:
            section_map[c] = s

    filled = 0
    for table in doc.tables:
        if len(table.rows) < 4:
            continue
        # Detect rating/evidence columns from header
        header = [c.text.strip().lower() for c in table.rows[0].cells]
        rate_col = None
        ev_col = None
        clause_col = 0
        for i, h in enumerate(header):
            if "rate" in h:
                rate_col = i
            elif "evidence" in h or "objective evid" in h:
                ev_col = i
            elif "ref" in h or "clause" in h:
                clause_col = i
        if rate_col is None:
            continue

        for row in table.rows[1:]:
            if clause_col >= len(row.cells):
                continue
            clause_text = row.cells[clause_col].text.strip()
            if not clause_text:
                continue
            # Match by prefix (e.g. "4.1" or "27001 4.1")
            match = None
            for c, s in section_map.items():
                if clause_text.startswith(c) or (c and clause_text.split()[-1].startswith(c)):
                    match = s
                    break
            if not match:
                continue
            status = match.get("status", "")
            if rate_col < len(row.cells) and status:
                row.cells[rate_col].text = str(rating_map.get(status, status))
            if ev_col and ev_col < len(row.cells):
                ev = match.get("evidence", "")
                if ev:
                    row.cells[ev_col].text = str(ev)[:500]
            filled += 1
    return filled


def fill_findings_table(doc, findings: list[dict], header_kw="no."):
    """Append findings/NC rows to a table whose header matches keywords."""
    table = find_table_by_header(doc, [header_kw, "description", "clause"])
    if not table or not findings:
        return 0
    added = 0
    for i, f in enumerate(findings):
        row = table.add_row()
        cells = row.cells
        values = [
            str(f.get("no", f.get("finding_id", i + 1))),
            f.get("description", ""),
            f.get("clause", ""),
        ]
        for j, v in enumerate(values):
            if j < len(cells):
                cells[j].text = ""
                p = cells[j].paragraphs[0]
                run = p.add_run(str(v))
                run.font.size = Pt(9)
        added += 1
    return added


def fill_template(doc_type: str, data: dict, template_path: str, out_path: str):
    """Fill a TÜV DOCX template with data; save to out_path. Returns out_path."""
    doc = Document(template_path)
    fill_info_block(doc, data)

    # Clause checklist sections
    sections = data.get("sections", [])
    if sections:
        fill_clause_ratings(doc, sections)

    # Findings / NCs
    findings = data.get("nonconformities") or data.get("findings") or []
    if findings:
        fill_findings_table(doc, findings)

    doc.save(out_path)
    return out_path