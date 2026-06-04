import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
TEMPLATES_DIR = os.path.join(BASE_DIR, 'templates', '04_03_26_consense_audit_documentation')

TUV_BLUE = RGBColor(0x00, 0x3D, 0x7A)
TUV_RED = RGBColor(0xC0, 0x00, 0x00)

TEMPLATE_MAP = {
    'Audit_Plan_Stage_1': 'Form \'FM-TAGMBH-MSZ-001_Auditplan-EN\' [English] (2).docx',
    'Audit_Plan_Stage_2': 'Form \'FM-TAGMBH-MSZ-001_Auditplan-EN\' [English] (2).docx',
    'Participation_List': 'Form \'FM-TAGMBH-MSZ-005_Participation-List-EN\' [English].docx',
    'Audit_Report': 'Form \'FM-TAGMBH-MSZ-003_Audit-Report-IMS-EN\' [English] (1).docx',
    'Certificate_Text': 'Form \'FM-TAGMBH-MSZ-038_Certificate-Text-EN\' [English].docx',
}

CHECKLIST_TEMPLATES = {
    'iso_9001': 'Form \'FM-TAGMBH-MSZ-010_Audit-Checklist-ISO-9001-EN\' [English] (1).docx',
    'iso_14001': 'Form \'FM-TAGMBH-MSZ-014_Audit-checklist-ISO-14001-EN\' [English] (1).docx',
    'iso_22000': 'Form \'FM-TAGMBH-MSZ-030_Audit-Checklist-ISO-22000-EN\' [English].docx',
    'iso_27001': 'Form \'FM-TAGMBH-MSZ-023_Audit-Checklist-ISO-27001-EN.xlsx',
    'iso_45001': 'Form \'FM-TAGMBH-MSZ-033-Audit-Checklist-combined-QM-EM-HSE-EN-EN\' [English] (1).docx',
    'iso_50001': 'Form \'FM-TAGMBH-MSZ-033-Audit-Checklist-combined-QM-EM-HSE-EN-EN\' [English] (1).docx',
    'iso_20000': 'ENG-Form \'FM-BA-ZET-MS-All_AQC_EN\' [English].docx',
    'iso_22301': 'BSO_Audit_Questionaire_ISO22301.docx',
}


def get_template_path(doc_type, standard_key=None):
    if doc_type == 'ISO_Checklist' and standard_key:
        fname = CHECKLIST_TEMPLATES.get(standard_key)
        if fname:
            path = os.path.join(TEMPLATES_DIR, fname)
            if os.path.exists(path):
                return path
    fname = TEMPLATE_MAP.get(doc_type)
    if not fname:
        return None
    path = os.path.join(TEMPLATES_DIR, fname)
    return path if os.path.exists(path) else None


def get_checklist_is_excel(standard_key):
    return standard_key == 'iso_27001'


def open_template(doc_type, standard_key=None):
    path = get_template_path(doc_type, standard_key)
    if not path:
        return None, None
    if path.endswith('.xlsx'):
        return 'xlsx', path
    return 'docx', Document(path)


def fill_cell(table, row_idx, col_idx, text, bold=False, color=None, size=9):
    if row_idx >= len(table.rows):
        return
    row = table.rows[row_idx]
    if col_idx >= len(row.cells):
        return
    cell = row.cells[col_idx]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def find_table_by_header(doc, keywords, min_rows=2):
    for table in doc.tables:
        if len(table.rows) < min_rows:
            continue
        header_text = ''
        for cell in table.rows[0].cells:
            header_text += cell.text.lower() + ' '
        if any(kw.lower() in header_text for kw in keywords):
            return table
    return None


def find_table_by_col_count(doc, cols, min_rows=2):
    for table in doc.tables:
        if len(table.rows) >= min_rows and len(table.rows[0].cells) == cols:
            return table
    return None


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def inject_checkbox_status(cell, status):
    checkbox = '☐'
    color = None
    if status == 'Conformant':
        checkbox = '☒'
    elif status == 'Non-Conformant':
        checkbox = '☒'
        color = TUV_RED
    elif status == 'Partially Conformant':
        checkbox = '☒'
        color = RGBColor(0xCC, 0x7A, 0x00)
    return f'{checkbox} {status}', color
