import os
import re
from docx import Document


def parse_docx(filepath):
    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)
    return {
        'filename': os.path.basename(filepath),
        'paragraphs': paragraphs,
        'text': '\n'.join(paragraphs),
        'tables': tables,
    }


def parse_txt(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
        text = f.read()
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    return {
        'filename': os.path.basename(filepath),
        'paragraphs': lines,
        'text': text,
        'tables': [],
    }


def extract_audit_notes(filepath):
    if filepath.endswith('.txt'):
        return parse_txt(filepath)
    return parse_docx(filepath)


def extract_manday_data(filepath):
    data = parse_docx(filepath)
    full_text = data['text']

    fields = {}

    client_match = re.search(r'(?:Client|Customer|Company|Organization)\s*[:;]\s*(.+?)(?:\n|$)', full_text, re.IGNORECASE)
    if client_match:
        fields['client_name'] = client_match.group(1).strip()

    address_match = re.search(r'(?:Address|Location|Site)\s*[:;]\s*(.+?)(?:\n|$)', full_text, re.IGNORECASE)
    if address_match:
        fields['client_address'] = address_match.group(1).strip()

    date_match = re.search(r'(?:Audit\s*Date|Date[s]?)\s*[:;]\s*(.+?)(?:\n|$)', full_text, re.IGNORECASE)
    if date_match:
        fields['audit_date'] = date_match.group(1).strip()

    manday_match = re.search(r'(\d+(?:\.\d+)?)\s*(?:Manday|Man[- ]?Day|MD|Audit\s*Day)', full_text, re.IGNORECASE)
    if manday_match:
        fields['mandays'] = manday_match.group(1)

    standard_match = re.search(r'(?:ISO|Standard)\s*[:;]\s*(.+?)(?:\n|$)', full_text, re.IGNORECASE)
    if standard_match:
        fields['standard'] = standard_match.group(1).strip()

    experts = []
    expert_pattern = re.compile(r'(?:Auditor|Expert|Assessor|Team\s*Member)\s*[:;]\s*(.+?)(?:\n|$)', re.IGNORECASE)
    for m in expert_pattern.finditer(full_text):
        experts.append(m.group(1).strip())
    if experts:
        fields['experts'] = experts

    for table in data['tables']:
        for row in table:
            row_lower = [c.lower() for c in row]
            if any('manday' in c or 'man-day' in c or 'man day' in c or 'audit days' in c for c in row_lower):
                fields['manday_table_row'] = ' | '.join(row)

    data['extracted'] = fields
    return data


def extract_manday_tables(docx_filepath):
    doc = Document(docx_filepath)
    tables = doc.tables

    result = {
        'employee_count': None,
        'equivalent_personnel': None,
        'base_mandays': None,
        'standards_text': [],
        'audit_type': None,
        'audit_dates': {},
        'team': [],
        'ims_systems': [],
        'ims_total': None,
        'ims_reduction_pct': None,
    }

    for table in tables:
        rows_text = []
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells]
            unique = list(dict.fromkeys(cells_text))
            row_key = ' | '.join(unique)
            rows_text.append(row_key)
            row_lower = row_key.lower()

            if 'total number of personnel' in row_lower:
                for cell in cells_text:
                    m = re.search(r'\b(\d+)\b', cell)
                    if m:
                        val = int(m.group(1))
                        if 0 < val < 100000:
                            result['employee_count'] = val

            if 'equivalent number of personnel' in row_lower and result['equivalent_personnel'] is None:
                for cell in cells_text:
                    m = re.search(r'\b(\d+)\b', cell)
                    if m:
                        val = int(m.group(1))
                        if 0 < val < 100000:
                            result['equivalent_personnel'] = val

            if 'final md calculation' in row_lower or 'final mandays calculation' in row_lower:
                for cell in cells_text:
                    m = re.search(r'(\d+(?:[,.]\s*\d+)?)\s*M\s*D', cell, re.I)
                    if m:
                        val = float(m.group(1).replace(',', '.'))
                        if result['base_mandays'] is None:
                            result['base_mandays'] = val

            dm = re.search(r'stage 1\s*=\s*(\d+(?:\.\d+)?)\s*MD', row_lower, re.I)
            if dm:
                result['audit_dates']['stage_1_md'] = float(dm.group(1))
            dm2 = re.search(r'stage 2\s*=\s*(\d+(?:\.\d+)?)\s*MD', row_lower, re.I)
            if dm2:
                result['audit_dates']['stage_2_md'] = float(dm2.group(1))

            if re.search(r'\bsurveillance\b', row_lower) and 'surveillance' not in str(result.get('audit_type', '')):
                result['audit_type'] = 'surveillance_1'
            if re.search(r'\brecertification\b', row_lower) and 'recertification' not in str(result.get('audit_type', '')):
                result['audit_type'] = 'recertification'
            if re.search(r'\binitial\b', row_lower) and result.get('audit_type') is None:
                result['audit_type'] = 'initial'

            if row_lower.startswith('standard ') and 'covered standards' not in row_lower:
                for cell in cells_text:
                    cell = cell.strip()
                    if cell and cell not in (
                        'Covered Standards', 'Standard 1', 'Standard 2',
                        'Standard 3', 'Standard 4', 'Other', ''
                    ):
                        if cell not in result['standards_text']:
                            result['standards_text'].append(cell)

            if row_lower.startswith('system ') and ':' in row_lower and 'level of integration' not in row_lower:
                for cell in cells_text:
                    m = re.match(r'System\s*(\d+):\s*(.+)', cell, re.I)
                    if m:
                        sys_num = int(m.group(1))
                        sys_name = m.group(2).strip()
                        while len(result['ims_systems']) < sys_num:
                            result['ims_systems'].append({})
                        result['ims_systems'][sys_num - 1] = {
                            'name': sys_name,
                            'mandays': result['ims_systems'][sys_num - 1].get('mandays'),
                        }

            if 'man-days according to relevant table' in row_lower:
                for cell in cells_text:
                    m = re.search(r'(\d+(?:[,.]\d+)?)\s*M\s*D', cell, re.I)
                    if m:
                        val = float(m.group(1).replace(',', '.'))
                        for sys in reversed(result['ims_systems']):
                            if sys.get('mandays') is None:
                                sys['mandays'] = val
                                break

            if 'total' in row_lower[:10] and ('md' in row_lower or 'md' in row_key):
                for cell in cells_text:
                    m = re.search(r'(\d+(?:[,.]\s*\d+)?)\s*M\s*D', cell, re.I)
                    if m:
                        val = float(m.group(1).replace(',', '').replace(' ', ''))
                        if result['ims_total'] is None:
                            result['ims_total'] = val

            pct_m = re.search(r'level of integration[^%]*?(\d+)\s*%', row_lower)
            if pct_m:
                pct = int(pct_m.group(1))
                if 0 < pct <= 50:
                    result['ims_reduction_pct'] = pct

    for table in tables:
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells]
            unique = list(dict.fromkeys(cells_text))
            row_key = ' | '.join(unique)
            row_lower = row_key.lower()

            if row_lower.startswith('la\t') or row_lower.startswith('la '):
                for cell in cells_text:
                    cell = cell.strip()
                    if cell not in ('LA', 'CA', 'TR', '') and len(cell) > 3:
                        if not any(m['role'] == 'Lead Auditor' for m in result['team']):
                            result['team'].append({'role': 'Lead Auditor', 'name': cell})

            if row_lower.startswith('ca\t') or row_lower.startswith('ca '):
                for cell in cells_text:
                    cell = cell.strip()
                    if cell not in ('LA', 'CA', 'TR', '') and len(cell) > 3:
                        if not any(m['name'] == cell for m in result['team']):
                            result['team'].append({'role': 'Co-Auditor', 'name': cell})

    if result['employee_count'] is None and result['equivalent_personnel'] is not None:
        result['employee_count'] = result['equivalent_personnel']

    return result
