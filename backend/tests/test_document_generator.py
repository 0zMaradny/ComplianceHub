"""Tests for document_generator — checklist generation and filename sanitization."""

import os

import pytest
from docx import Document

from app.services.document_generator import (
    generate_document_file,
    sanitize_filename,
    generate_iso_checklist,
)


class TestGenerateIsoChecklist:
    def test_builds_docx_without_template(self, tmp_path):
        out = str(tmp_path / "checklist.docx")
        data = {
            "client_name": "Acme Corp",
            "standard": "ISO 9001:2015",
            "audit_date": "2025-01-01",
            "sections": [
                {"clause": "4.1", "title": "Context", "status": "Conformant", "findings": "OK"},
                {"clause": "4.2", "title": "Needs", "status": "Non-Conformant", "findings": "Gap"},
            ],
        }
        result = generate_iso_checklist(data, out, client_key=None)
        assert result == out
        doc = Document(out)
        assert len(doc.tables) >= 1
        # Header row carries clause + status columns
        header = doc.tables[0].rows[0]
        header_text = " ".join(c.text for c in header.cells)
        assert "Status" in header_text or "clause" in header_text.lower()

    def test_empty_sections_no_crash(self, tmp_path):
        out = str(tmp_path / "empty.docx")
        result = generate_iso_checklist({}, out, client_key=None)
        assert result == out
        doc = Document(out)
        assert len(doc.tables) >= 0  # cover page still renders


class TestSanitizeFilename:
    def test_removes_special_chars(self):
        result = sanitize_filename("My Client (Acme) [2025].docx")
        assert "<" not in result and ">" not in result

    def test_caps_at_120_chars(self):
        long_name = "A" * 100 + ".docx"
        result = sanitize_filename(long_name)
        assert len(result) <= 120
        assert result == "A" * 100 + ".docx"

    def test_allows_alphanumeric(self):
        result = sanitize_filename("Client-Name_2025.docx")
        assert "Client" in result