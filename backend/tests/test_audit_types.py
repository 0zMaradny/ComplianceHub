"""Tests for audit-type-aware document generation."""
import os, tempfile
import pytest
from docx import Document

from app.services.audit_types import get_docs_for_audit_type, get_hard_rules, get_audit_type_label, SURVEILLANCE_DOCS, FULL_CERTIFICATION_DOCS
from app.services.doc_schemas import build_document_data, DOCUMENT_CONTRACTS
from app.services.document_generator import generate_surveillance_plan, generate_surveillance_report, generate_surveillance_findings_summary


def test_initial_full_package():
    assert get_docs_for_audit_type("initial") == FULL_CERTIFICATION_DOCS
    assert len(get_docs_for_audit_type("initial")) == 8


def test_recert_full_package():
    assert get_docs_for_audit_type("recertification") == FULL_CERTIFICATION_DOCS
    assert len(get_docs_for_audit_type("recertification")) == 8


def test_surveillance_reduced():
    docs = get_docs_for_audit_type("surveillance")
    assert docs == SURVEILLANCE_DOCS and len(docs) == 5
    assert "Certificate" not in docs and "Certificate_Text" not in docs
    assert "Surveillance_Plan" in docs
    assert "Audit_Plan_Stage_1" not in docs and "Audit_Plan_Stage_2" not in docs


def test_unknown_fallback():
    assert get_docs_for_audit_type("transfer") == FULL_CERTIFICATION_DOCS


def test_hard_rules_all_types():
    for at in ("initial", "surveillance", "recertification"):
        assert get_hard_rules(at)


def test_surv_rules_content():
    r = get_hard_rules("surveillance")
    assert "Certificate" in r and "Review of Open Nonconformities" in r
    assert "Review of Changes" in r and "Continued Certification" in r


def test_recert_rules_content():
    r = get_hard_rules("recertification")
    assert "3-Year Performance Review" in r and "renewal" in r.lower()


def test_unknown_rules_empty():
    assert get_hard_rules("transfer") == ""


def test_labels():
    assert get_audit_type_label("initial") == "Initial Certification"
    assert get_audit_type_label("surveillance") == "Surveillance"
    assert get_audit_type_label("recertification") == "Recertification"


@pytest.mark.parametrize("dt", ["Surveillance_Plan", "Surveillance_Report", "Surveillance_Findings_Summary"])
def test_contracts_exist(dt):
    assert dt in DOCUMENT_CONTRACTS


def test_build_surv_plan():
    d = build_document_data("Surveillance_Plan", {})
    assert d["client_name"] == "" and isinstance(d["audit_team"], list) and isinstance(d["daily_schedule"], list)


def test_build_surv_report():
    d = build_document_data("Surveillance_Report", {"cb_recommendation": "Continued Certification"})
    assert d["cb_recommendation"] == "Continued Certification"
    assert "previously_raised_ncs" in d and "changes_since_initial" in d


def test_build_findings():
    d = build_document_data("Surveillance_Findings_Summary", {})
    assert isinstance(d["findings"], list) and d["summary"]["total_findings"] == ""


def _plan_data():
    return {
        "client_name": "Test Client", "audit_date": "01/01/2026", "standard": "ISO 9001:2015",
        "surveillance_cycle": "Year 1 Surveillance",
        "audit_team": [{"name": "A. Auditor", "role": "Lead", "days": 2}],
        "audit_objectives": ["Review of open NCs", "Review of changes"],
        "audit_scope": "Reduced scope.", "audit_criteria": ["clauses 4-10"],
        "review_of_open_ncs": "All closed.", "review_of_changes": "New ERP.",
        "daily_schedule": [{"day": 1, "date": "01/01/2026", "time": "09:00", "activity": "Opening", "auditee": "Mgmt", "auditor": "A. Auditor", "clause": "4.1"}],
        "confidentiality": "Conf", "language": "English", "report_date": "31/01/2026",
    }


def test_gen_plan_docx():
    with tempfile.TemporaryDirectory() as tmp:
        p = generate_surveillance_plan(_plan_data(), os.path.join(tmp, "p.docx"))
        assert os.path.exists(p)
        t = "\n".join(x.text for x in Document(p).paragraphs)
        assert "Surveillance Audit Plan" in t and "Review of Open Nonconformities" in t and "Review of Changes" in t


def test_gen_report_docx():
    with tempfile.TemporaryDirectory() as tmp:
        d = _plan_data()
        d.update({"report_number": "TUV-SR-2026-001", "previously_raised_ncs": "Reviewed.",
                  "changes_since_initial": "Noted.", "cb_recommendation": "Continued Certification",
                  "findings_summary": "Effective.", "conclusion": "Continues."})
        p = generate_surveillance_report(d, os.path.join(tmp, "r.docx"))
        assert os.path.exists(p)
        t = "\n".join(x.text for x in Document(p).paragraphs)
        assert "Previous" in t and "Changes Since Initial" in t and "CB Recommendation" in t


def test_gen_findings_docx():
    with tempfile.TemporaryDirectory() as tmp:
        d = {"client_name": "Test Client", "audit_date": "01/01/2026", "standard": "ISO 9001:2015",
             "cycle": "Year 1", "summary": {"total_findings": 1, "major": 0, "minor": 1, "ofi": 0, "observations": 0, "recurring_ncs": 0},
             "findings": [{"finding_id": "SF-001", "clause": "8.5", "type": "NC", "severity": "Minor",
                           "description": "Minor issue.", "status": "Open", "previous_nc_id": ""}],
             "overall_assessment": "OK."}
        p = generate_surveillance_findings_summary(d, os.path.join(tmp, "f.docx"))
        assert os.path.exists(p)
        doc = Document(p)
        t = "\n".join(x.text for x in doc.paragraphs)
        table_text = "\n".join(c.text for tb in doc.tables for row in tb.rows for c in row.cells)
        assert "Surveillance Findings Summary" in t and "SF-001" in table_text


def test_pipeline_targets():
    from app.services.audit_types import AUDIT_TYPE_DOCS
    assert set(AUDIT_TYPE_DOCS["surveillance"]) == set(SURVEILLANCE_DOCS)
    assert "Certificate" not in AUDIT_TYPE_DOCS["surveillance"]
    assert len(AUDIT_TYPE_DOCS["initial"]) == len(AUDIT_TYPE_DOCS["recertification"]) == 8