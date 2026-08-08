"""E2E: verify TUV template fill preserves layout + injects data."""
import sys, os, tempfile
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from docx import Document
from app.services.document_generator import generate_document_file

d = {"client_name": "Test Corp", "standard": "ISO 9001:2015",
     "audit_date": "01/01/2026", "lead_auditor": "Osama El Maradny",
     "scope": "Quality management", "audit_type": "CERTIFICATION"}

for doc_type, std_key in [("Audit_Plan_Stage_1", "iso_9001"),
                          ("Audit_Report", "iso_9001"),
                          ("Participation_List", "iso_9001")]:
    out = tempfile.mkdtemp()
    p = generate_document_file(doc_type, d, out, standard_key=std_key)
    doc = Document(p)
    txt = "\n".join(x.text for x in doc.paragraphs)
    tbl = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    print(f"{doc_type}: out={os.path.basename(p)} layout={'Company information' in tbl or 'Company name' in tbl} client={'Test Corp' in tbl or 'Test Corp' in txt}")