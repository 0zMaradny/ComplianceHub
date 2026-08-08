"""Generate the FIXED TÜV AUSTRIA form templates as structural placeholders.

Run: ``python scripts/generate_templates.py``

The proprietary TÜV binaries were not recoverable after the environment reset, so
this regenerates *structurally faithful* forms (labelled info-block cells +
clause-rating tables + findings tables) with a real ``word/webSettings.xml``
so the ``clean_docx`` bug/fix is verifiable. The filler is label-/token-driven,
so the real TÜV binaries can replace these 1:1 with NO code changes.
"""

from __future__ import annotations

import os
import re
import shutil
import zipfile

from docx import Document

try:
    from openpyxl import Workbook
except Exception:  # pragma: no cover
    Workbook = None

BACKEND_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATES_ROOT = os.path.join(BACKEND_DIR, "templates", "04_03_26_consense_audit_documentation")

# Filenames MUST mirror app.services.template_manager.TEMPLATE_FILES exactly.
FORM_FN = {
    "MSZ-001_AuditPlan": "Form FM-TAGMBH-MSZ-001_Auditplan-EN.docx",
    "MSZ-002_ISMSPlan": "Form 'FM-TAGMBH-MSZ-002_Auditplan-ISMS-EN' (1).docx",
    "MSZ-003_AuditReport": "Form FM-TAGMBH-MSZ-003_Auditreport-EN.docx",
    "MSZ-005_ParticipationList": "Form FM-TAGMBH-MSZ-005_ParticipationList-EN.docx",
    "MSZ-033_Checklist": "Form FM-TAGMBH-MSZ-033_CombinedChecklist-EN.docx",
    "MSZ-023_27001": "Form FM-TAGMBH-MSZ-023_Checklist27001-EN.xlsx",
    "MSZ-038_CertText": "Form FM-TAGMBH-MSZ-038_CertText-EN.docx",
    "BSO22301_Questionnaire": "Form FM-TAGMBH-BSO22301_Questionnaire-EN.docx",
}

WEBSETTINGS_PATH = "word/webSettings.xml"
WEBSETTINGS_RELTYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/webSettings"
)

INFO_LABELS = [
    "Company name", "Client", "Full name", "Address", "Contact person", "Contact",
    "Email", "Phone", "Website", "Audit type", "Standard", "Scope",
    "Audit period", "Audit dates", "Auditor", "Auditor firm", "Audit stage", "Stage",
]

PLAN_CLAUSES = [
    ("4.1", "Understanding the organisation"),
    ("4.2", "Understanding needs & expectations"),
    ("5.2", "Leadership - policy"),
    ("8.2", "Customer requirements"),
    ("8.3", "Design & development"),
    ("SF-001", "ISMS control A.8 - asset ownership"),
]
ISMS_CLAUSES = [
    ("5.1", "Policies"), ("5.2", "Risk acceptance"), ("8.2", "Customer reqs"),
    ("8.3", "Design"), ("SF-001", "ISMS control A.8"),
]


def _ensure_websettings(path: str) -> None:
    """Guarantee the package carries word/webSettings.xml + its relationship."""
    if not os.path.exists(path) or not path.endswith(".docx"):
        return
    with zipfile.ZipFile(path, "r") as z:
        names = set(z.namelist())
        members = {n: z.read(n) for n in names}
    changed = False
    if WEBSETTINGS_PATH not in names:
        members[WEBSETTINGS_PATH] = (
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
            b'<w:webSettings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            b' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" mc:Ignorable="w14 w15 w16">'
            b'<w:compat><w:compatSetting w:name="compatMode" w:val="1"/></w:compat>'
            b'</w:webSettings>'
        )
        changed = True
    rels_name = "word/_rels/document.xml.rels"
    if rels_name in members:
        rels = members[rels_name].decode("utf-8", errors="replace")
        if WEBSETTINGS_RELTYPE not in rels:
            new_rel = (
                '  <Relationship Id="rIdWebSettings" '
                f'Type="{WEBSETTINGS_RELTYPE}" Target="{WEBSETTINGS_PATH}" />\n'
            )
            rels = re.sub(r"</Relationships>", new_rel + "</Relationships>", rels, count=1)
            members[rels_name] = rels.encode("utf-8")
            changed = True
    if changed:
        tmp = path + ".gen.tmp"
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as out:
            for n, data in members.items():
                out.writestr(n, data)
        shutil.move(tmp, path)


def _add_info_block(doc: Document) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label in INFO_LABELS:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = ""


def _add_rating_table(doc: Document, clauses: list) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text, h[1].text, h[2].text = "Clause", "Control", "Rating"
    for ref, title in clauses:
        row = table.add_row()
        row.cells[0].text = ref
        row.cells[1].text = title
        row.cells[2].text = ""


def _add_findings_table(doc: Document) -> None:
    cols = ["Finding", "Clause", "Area", "Description", "Severity", "Cause", "Status"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, h in enumerate(cols):
        table.rows[0].cells[i].text = h


def _make_plan(isms: bool) -> Document:
    doc = Document()
    doc.add_heading("Audit Plan", level=0)
    _add_info_block(doc)
    doc.add_paragraph("")
    doc.add_heading("Checklist", level=1)
    _add_rating_table(doc, ISMS_CLAUSES if isms else PLAN_CLAUSES)
    doc.add_paragraph("")
    doc.add_heading("Non-conformities", level=1)
    _add_findings_table(doc)
    return doc


def _make_report() -> Document:
    doc = Document()
    doc.add_heading("Audit Report", level=0)
    _add_info_block(doc)
    doc.add_paragraph("")
    doc.add_heading("Findings", level=1)
    _add_findings_table(doc)
    doc.add_paragraph("")
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("Conclusion text to be filled.")
    return doc


def _make_participation() -> Document:
    doc = Document()
    doc.add_heading("Participation List", level=0)
    _add_info_block(doc)
    doc.add_paragraph("")
    doc.add_heading("Participants", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["Name", "Organisation", "Role", "Signature", "Date"]):
        table.rows[0].cells[i].text = h
    return doc


def _make_checklist() -> Document:
    doc = Document()
    doc.add_heading("Combined Audit Checklist", level=0)
    _add_info_block(doc)
    doc.add_paragraph("")
    _add_rating_table(doc, PLAN_CLAUSES)
    return doc


def _make_cert() -> Document:
    doc = Document()
    doc.add_heading("Certificate Text", level=0)
    _add_info_block(doc)
    doc.add_paragraph("")
    doc.add_paragraph("This is to certify that «Company name» ... against «Standard».")
    doc.add_paragraph("Scope: «Scope»  |  Period: «Audit period»")
    return doc


def _make_bso() -> Document:
    doc = Document()
    doc.add_heading("BSO 22301 Questionnaire", level=0)
    _add_info_block(doc)
    doc.add_paragraph("")
    doc.add_heading("Questions", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text, table.rows[0].cells[1].text = "Question", "Answer"
    for q in ["Q1. Is there a BCMS?", "Q2. Are roles defined?", "Q3. Is there a testing plan?"]:
        row = table.add_row()
        row.cells[0].text = q
        row.cells[1].text = ""
    return doc


def _make_msz023():
    if Workbook is None:
        return None
    wb = Workbook()
    ws = wb.active
    ws.title = "ISMS 27001 Checklist"
    ws.append(["Clause", "Control", "Rating"])
    for ref, title in ISMS_CLAUSES:
        ws.append([ref, title, ""])
    ws.append(["Company name", "Test Corp", "", "", "", "", ""])
    return wb


def generate_all() -> int:
    os.makedirs(TEMPLATES_ROOT, exist_ok=True)
    targets = {
        "MSZ-001_AuditPlan": ("docx", False),
        "MSZ-002_ISMSPlan": ("docx", True),
        "MSZ-003_AuditReport": ("docx", None),
        "MSZ-005_ParticipationList": ("docx", None),
        "MSZ-033_Checklist": ("docx", None),
        "MSZ-023_27001": ("xlsx", None),
        "MSZ-038_CertText": ("docx", None),
        "BSO22301_Questionnaire": ("docx", None),
    }
    builders = {
        "MSZ-001_AuditPlan": _make_plan,
        "MSZ-002_ISMSPlan": _make_plan,
        "MSZ-003_AuditReport": _make_report,
        "MSZ-005_ParticipationList": _make_participation,
        "MSZ-033_Checklist": _make_checklist,
        "MSZ-038_CertText": _make_cert,
        "BSO22301_Questionnaire": _make_bso,
    }
    built = 0
    for key, (kind, isms) in targets.items():
        fn = FORM_FN[key]
        path = os.path.join(TEMPLATES_ROOT, fn)
        if kind == "docx":
            builder = builders[key]
            isms_val = isms if isms is not None else False
            doc = builder(isms_val) if key in ("MSZ-001_AuditPlan", "MSZ-002_ISMSPlan") else builder()
            doc.save(path)
            _ensure_websettings(path)
        else:
            wb = _make_msz023()
            wb.save(path)
        built += 1
        print(f"  + {fn}")
    print(f"generated {built} templates into {TEMPLATES_ROOT}")
    return built


if __name__ == "__main__":
    raise SystemExit(0 if generate_all() else 1)
